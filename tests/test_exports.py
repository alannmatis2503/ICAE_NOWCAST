"""Test complet des exports Word et Excel avec donnees reelles."""
import sys
import os
os.environ["PYTHONIOENCODING"] = "utf-8"
sys.path.insert(0, r"c:\Users\HP\Documents\Stage pro BEAC\Work\ICAE\Mars 2026\Apps\ICAE_Streamlit")

import pandas as pd
import numpy as np
from pathlib import Path

from config import CONSOLIDES, COUNTRY_CODES, COUNTRY_NAMES, POIDS_PIB, LOGO_PATH
from core.icae_engine import run_icae_pipeline
from core.quarterly import quarterly_mean, calc_ga_trim, calc_gt_trim
from core.cemac_engine import compute_icae_cemac, quarterly_cemac
from core.forecast_engine import run_all_forecasts
from core.nowcast_engine import run_nowcast, compute_ins_out_metrics
from io_utils.excel_reader import (
    read_consignes, read_codification, read_donnees_calcul, rename_columns_to_codes,
)
from io_utils.excel_writer import (
    write_icae_output, write_previsions_excel,
    write_nowcast_excel, write_cemac_excel,
)
from io_utils.word_report import generate_note_icae, generate_note_nowcast
from ui.charts import chart_ga_bars, chart_icae_monthly, fig_to_png_bytes

OUT_DIR = Path(r"c:\Users\HP\Documents\Stage pro BEAC\Work\ICAE\Mars 2026\Apps\ICAE_Streamlit\tests\output")
OUT_DIR.mkdir(parents=True, exist_ok=True)

errors = []

def assert_file(path, min_size=500):
    """Verifie qu'un fichier existe et a une taille minimale."""
    if not path.exists():
        raise AssertionError(f"Fichier non cree: {path}")
    size = path.stat().st_size
    if size < min_size:
        raise AssertionError(f"Fichier trop petit ({size} bytes): {path}")
    return size


# ====================================================================
# PREPARATION : charger CMR et lancer le pipeline
# ====================================================================
print("=" * 60)
print("PREPARATION : Pipeline ICAE pour CMR")
print("=" * 60)

test_file = CONSOLIDES / "ICAE_CMR_Consolide.xlsx"
consignes = read_consignes(test_file)
codification = read_codification(test_file)
donnees = read_donnees_calcul(test_file)
donnees = rename_columns_to_codes(donnees, codification)

base_year = consignes.get("base_year", 2023)
dates = pd.to_datetime(donnees["Date"])
base_mask = dates.dt.year == base_year
base_indices = donnees.index[base_mask]
base_rows = range(base_indices[0], base_indices[-1] + 1) if len(base_indices) > 0 else range(108, 120)

priors = pd.Series(
    codification["PRIOR"].values,
    index=codification["Code"].values,
    dtype=float,
).fillna(0)

results = run_icae_pipeline(donnees, priors, base_year, base_rows)
q = quarterly_mean(results["icae"], results["dates"])
q["GA_Trim"] = calc_ga_trim(q["icae_trim"])
q["GT_Trim"] = calc_gt_trim(q["icae_trim"])
# Convertir Period -> str pour serialisation
q["quarter"] = q["quarter"].astype(str)
q["debut"] = q["debut"].astype(str)
q["fin"] = q["fin"].astype(str)

print(f"  Pipeline OK : ICAE {results['icae'].dropna().shape[0]} obs, last={results['icae'].dropna().iloc[-1]:.2f}")

# ====================================================================
# TEST 1 : Export Excel ICAE (_OUT.xlsx)
# ====================================================================
print("\n" + "=" * 60)
print("TEST 1 : Export Excel ICAE (_OUT.xlsx)")
print("=" * 60)

try:
    export_results = {
        "icae": results["icae"],
        "indice": results["indice"],
        "sum_m": results["sum_m"],
        "pond_finale": results["pond_finale"],
        "quarterly": q,
    }
    data = write_icae_output(test_file, export_results, "CMR")
    out_path = OUT_DIR / "ICAE_CMR_OUT.xlsx"
    out_path.write_bytes(data)
    size = assert_file(out_path, min_size=5000)
    
    # Verifier le contenu
    import openpyxl
    wb = openpyxl.load_workbook(out_path, data_only=True)
    sheets = wb.sheetnames
    wb.close()
    print(f"  [OK] Fichier: {out_path.name} ({size:,} bytes)")
    print(f"       Feuilles: {sheets}")
except Exception as e:
    errors.append(f"Excel ICAE: {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# TEST 2 : Export Excel Previsions
# ====================================================================
print("\n" + "=" * 60)
print("TEST 2 : Export Excel Previsions")
print("=" * 60)

try:
    data_cols = [c for c in donnees.columns if c != "Date"]
    test_vars = data_cols[:3]  # 3 premieres variables
    
    all_prev_results = {}
    for var in test_vars:
        series = donnees[var].dropna()
        if len(series) >= 24:
            all_prev_results[var] = run_all_forecasts(
                series, h=3, methods=["MM3", "MM6", "NS", "TL"], bt_window=12
            )
    
    # Preparer les structures
    previsions = {}
    backtesting_data = {}
    recommandations = {}
    for var, res in all_prev_results.items():
        previsions[var] = res["forecasts"]
        backtesting_data[var] = res["backtesting"]
        best = res["best_method"]
        recommandations[var] = {
            "method": best,
            "mape": res["backtesting"].get(best, {}).get("mape", np.nan),
        }
    
    edited_df = pd.DataFrame({
        "Date": pd.date_range("2026-01-01", periods=3, freq="MS").strftime("%Y-%m"),
    })
    for var in test_vars:
        if var in all_prev_results:
            best = all_prev_results[var]["best_method"]
            fc = all_prev_results[var]["forecasts"][best]
            edited_df[var] = fc[:3] if len(fc) >= 3 else np.nan
    
    data = write_previsions_excel(
        donnees, previsions, backtesting_data, recommandations,
        edited=edited_df,
    )
    out_path = OUT_DIR / "Prevision_CMR_Q1_2026.xlsx"
    out_path.write_bytes(data)
    size = assert_file(out_path, min_size=2000)
    
    # Verifier
    wb = openpyxl.load_workbook(out_path, data_only=True)
    sheets = wb.sheetnames
    wb.close()
    print(f"  [OK] Fichier: {out_path.name} ({size:,} bytes)")
    print(f"       Feuilles: {sheets}")
    print(f"       Variables: {list(all_prev_results.keys())}")
except Exception as e:
    errors.append(f"Excel Previsions: {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# TEST 3 : Export Excel Nowcast
# ====================================================================
print("\n" + "=" * 60)
print("TEST 3 : Export Excel Nowcast")
print("=" * 60)

try:
    # Creer des donnees synthetiques pour PIB trimestriel (pas de fichier reel)
    n_q = 40
    pib_q_idx = pd.period_range("2015Q1", periods=n_q, freq="Q")
    np.random.seed(42)
    pib_values = 100 + np.cumsum(np.random.randn(n_q) * 2)
    pib_q = pd.Series(pib_values, index=pib_q_idx, name="PIB")
    
    # Creer des HF trimestriels synthetiques
    hf_q = pd.DataFrame(
        np.random.randn(n_q, 5) * 10 + 50,
        index=pib_q_idx,
        columns=[f"HF_{i}" for i in range(5)],
    )
    # Correlationner HF_0 au PIB
    hf_q["HF_0"] = pib_values * 0.8 + np.random.randn(n_q) * 5
    
    nowcast_results = run_nowcast(
        pib_q, hf_q,
        models=["Bridge", "U-MIDAS", "PC", "DFM"],
        h_ahead=4, n_components=2,
    )
    
    params = {"Pays": "CMR", "Modeles": "Bridge, U-MIDAS, PC, DFM",
              "Nb facteurs": 2, "Horizon": 4}
    data = write_nowcast_excel(pib_q, nowcast_results, params)
    out_path = OUT_DIR / "RESULT_NOWCAST_CMR.xlsx"
    out_path.write_bytes(data)
    size = assert_file(out_path, min_size=2000)
    
    wb = openpyxl.load_workbook(out_path, data_only=True)
    sheets = wb.sheetnames
    wb.close()
    print(f"  [OK] Fichier: {out_path.name} ({size:,} bytes)")
    print(f"       Feuilles: {sheets}")
    print(f"       Modeles: {list(nowcast_results.keys())}")
except Exception as e:
    errors.append(f"Excel Nowcast: {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# TEST 4 : Export Excel CEMAC
# ====================================================================
print("\n" + "=" * 60)
print("TEST 4 : Export Excel CEMAC")
print("=" * 60)

try:
    icae_dict = {}
    for code in COUNTRY_CODES:
        fpath = CONSOLIDES / f"ICAE_{code}_Consolide.xlsx"
        if not fpath.exists():
            continue
        cons = read_consignes(fpath)
        cod = read_codification(fpath)
        don = read_donnees_calcul(fpath)
        don = rename_columns_to_codes(don, cod)
        
        by = cons.get("base_year", 2023)
        dt = pd.to_datetime(don["Date"])
        bm = dt.dt.year == by
        bi = don.index[bm]
        br = range(bi[0], bi[-1] + 1) if len(bi) > 0 else range(108, 120)
        
        pr = pd.Series(cod["PRIOR"].values, index=cod["Code"].values, dtype=float).fillna(0)
        res = run_icae_pipeline(don, pr, by, br)
        icae_series = res["icae"]
        icae_series.index = dt
        icae_dict[code] = icae_series
    
    cemac_df = compute_icae_cemac(icae_dict, POIDS_PIB)
    dates_cemac = pd.to_datetime(cemac_df.index)
    q_cemac = quarterly_cemac(cemac_df, dates_cemac)
    
    # Convertir Period -> str pour serialisation
    if "Trimestre" in q_cemac.columns:
        q_cemac["Trimestre"] = q_cemac["Trimestre"].astype(str)
    
    data = write_cemac_excel(cemac_df, q_cemac, POIDS_PIB)
    out_path = OUT_DIR / "ICAE_CEMAC_OUT.xlsx"
    out_path.write_bytes(data)
    size = assert_file(out_path, min_size=2000)
    
    wb = openpyxl.load_workbook(out_path, data_only=True)
    sheets = wb.sheetnames
    wb.close()
    print(f"  [OK] Fichier: {out_path.name} ({size:,} bytes)")
    print(f"       Feuilles: {sheets}")
    print(f"       Pays: {list(icae_dict.keys())}")
except Exception as e:
    errors.append(f"Excel CEMAC: {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# TEST 5 : Export Word - Note ICAE (sans graphiques)
# ====================================================================
print("\n" + "=" * 60)
print("TEST 5 : Export Word - Note ICAE (sans graphiques)")
print("=" * 60)

try:
    user_texts = {
        "accroche_evolution": "Le rythme de progression des activites economiques a accelere au T4 2025.",
        "paragraphs_evolution": [
            "Au cours du quatrieme trimestre 2025, le secteur productif du Cameroun "
            "a montre des signes de reprise. L'ICAE du Cameroun s'est accru de 2,07 % "
            "en glissement annuel, apres 3,5 % au trimestre precedent.",
            "Les activites extractives (petrole) ont contribue positivement.",
            "L'industrie manufacturiere a progresse, notamment les telecoms.",
            "Le secteur des transports a soutenu la dynamique globale.",
        ],
        "accroche_perspectives": "Le rythme devrait se maintenir au T1 2026.",
        "paragraphs_perspectives": [
            "Les projections sectorielles indiquent une stabilite.",
            "Les facteurs de risque incluent les tensions sur les prix du petrole.",
            "En synthese, l'ICAE du Cameroun devrait progresser de 1,8 % en GA.",
        ],
    }
    
    logo = str(LOGO_PATH) if LOGO_PATH.exists() else None
    
    doc_bytes = generate_note_icae(
        country_name="Cameroun",
        trimestre="T4",
        annee=2025,
        icae_data={},
        chart_evolution_bytes=None,
        chart_perspectives_bytes=None,
        logo_path=logo,
        user_texts=user_texts,
    )
    
    out_path = OUT_DIR / "Note_ICAE_CMR_T4_2025_sans_graph.docx"
    out_path.write_bytes(doc_bytes)
    size = assert_file(out_path, min_size=3000)
    
    # Verifier le contenu Word
    from docx import Document
    doc = Document(str(out_path))
    n_paragraphs = len(doc.paragraphs)
    title_text = doc.paragraphs[1].text if len(doc.paragraphs) > 1 else ""
    print(f"  [OK] Fichier: {out_path.name} ({size:,} bytes)")
    print(f"       Paragraphes: {n_paragraphs}")
    print(f"       Titre: {title_text[:80]}...")
    has_logo = any(r.element.tag.endswith('}drawing') for p in doc.paragraphs for r in p.runs)
    print(f"       Logo inclus: {has_logo or 'image' in str([p.text for p in doc.paragraphs[:3]])}")
except Exception as e:
    errors.append(f"Word Note ICAE (sans graphique): {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# TEST 6 : Export Word - Note ICAE (avec graphiques)
# ====================================================================
print("\n" + "=" * 60)
print("TEST 6 : Export Word - Note ICAE (avec graphiques)")
print("=" * 60)

try:
    # Generer les graphiques
    ga = results["ga_monthly"]
    fig_evo = chart_ga_bars(
        results["dates"], ga.values,
        title="ICAE CMR - Glissement annuel (%)",
    )
    chart_evo_bytes = fig_to_png_bytes(fig_evo)
    
    fig_persp = chart_icae_monthly(
        results["dates"], results["icae"].values,
        title="ICAE CMR - Evolution mensuelle",
    )
    chart_persp_bytes = fig_to_png_bytes(fig_persp)
    
    if chart_evo_bytes is None:
        print("  [WARN] Kaleido n'a pas pu generer chart_evo_bytes, test avec image dummy")
        # Creer une petite image PNG dummy (1x1 pixel blanc)
        import struct, zlib
        def _make_dummy_png():
            raw = b'\x00\xFF\xFF\xFF'
            compressed = zlib.compress(raw)
            def chunk(tp, d):
                return struct.pack('>I', len(d)) + tp + d + struct.pack('>I', zlib.crc32(tp + d) & 0xFFFFFFFF)
            return (b'\x89PNG\r\n\x1a\n' +
                    chunk(b'IHDR', struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0)) +
                    chunk(b'IDAT', compressed) +
                    chunk(b'IEND', b''))
        chart_evo_bytes = _make_dummy_png()
        chart_persp_bytes = _make_dummy_png()
    
    doc_bytes = generate_note_icae(
        country_name="Cameroun",
        trimestre="T4",
        annee=2025,
        icae_data={},
        chart_evolution_bytes=chart_evo_bytes,
        chart_perspectives_bytes=chart_persp_bytes,
        logo_path=str(LOGO_PATH) if LOGO_PATH.exists() else None,
        user_texts=user_texts,
    )
    
    out_path = OUT_DIR / "Note_ICAE_CMR_T4_2025_avec_graph.docx"
    out_path.write_bytes(doc_bytes)
    size = assert_file(out_path, min_size=5000)
    
    doc = Document(str(out_path))
    n_paragraphs = len(doc.paragraphs)
    # Compter les images
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    n_images = len([r for r in doc.part.rels.values() if "image" in r.reltype])
    print(f"  [OK] Fichier: {out_path.name} ({size:,} bytes)")
    print(f"       Paragraphes: {n_paragraphs}")
    print(f"       Images embarquees: {n_images}")
except Exception as e:
    errors.append(f"Word Note ICAE (avec graphiques): {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# TEST 7 : Export Word - Note Nowcast
# ====================================================================
print("\n" + "=" * 60)
print("TEST 7 : Export Word - Note Nowcast")
print("=" * 60)

try:
    results_for_report = {}
    
    # Utiliser les resultats nowcast du test 3
    perf_rows = []
    for name, r in nowcast_results.items():
        m = r["metrics"]
        perf_rows.append({
            "Modele": name,
            "RMSE_in": round(m["in_sample"].get("rmse", np.nan), 2),
            "MAE_in": round(m["in_sample"].get("mae", np.nan), 2),
            "MAPE_in": round(m["in_sample"].get("mape", np.nan), 2),
            "RMSE_out": round(m["out_sample"].get("rmse", np.nan), 2),
            "MAE_out": round(m["out_sample"].get("mae", np.nan), 2),
            "MAPE_out": round(m["out_sample"].get("mape", np.nan), 2),
        })
    results_for_report["CMR"] = {
        "metrics_df": pd.DataFrame(perf_rows),
    }
    
    doc_bytes = generate_note_nowcast(
        results_by_country=results_for_report,
        chart_bytes={},
        logo_path=str(LOGO_PATH) if LOGO_PATH.exists() else None,
    )
    
    out_path = OUT_DIR / "Note_Nowcast_2025.docx"
    out_path.write_bytes(doc_bytes)
    size = assert_file(out_path, min_size=3000)
    
    doc = Document(str(out_path))
    n_paragraphs = len(doc.paragraphs)
    # Compter les tables
    n_tables = len(doc.tables)
    print(f"  [OK] Fichier: {out_path.name} ({size:,} bytes)")
    print(f"       Paragraphes: {n_paragraphs}")
    print(f"       Tables: {n_tables}")
    if n_tables > 0:
        t = doc.tables[0]
        print(f"       Table 1 : {len(t.rows)} lignes x {len(t.columns)} colonnes")
        # Afficher les en-tetes
        headers = [t.cell(0, j).text for j in range(len(t.columns))]
        print(f"       En-tetes: {headers}")
except Exception as e:
    errors.append(f"Word Note Nowcast: {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# TEST 8 : Validation des fichiers Excel generes (lecture)
# ====================================================================
print("\n" + "=" * 60)
print("TEST 8 : Validation contenu Excel")
print("=" * 60)

try:
    # Previsions
    prev_path = OUT_DIR / "Prevision_CMR_Q1_2026.xlsx"
    xls = pd.ExcelFile(prev_path)
    print(f"  Previsions ({prev_path.name}):")
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        print(f"    {sheet}: {df.shape[0]} lignes x {df.shape[1]} colonnes")
    xls.close()
    
    # Nowcast
    now_path = OUT_DIR / "RESULT_NOWCAST_CMR.xlsx"
    xls = pd.ExcelFile(now_path)
    print(f"  Nowcast ({now_path.name}):")
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        print(f"    {sheet}: {df.shape[0]} lignes x {df.shape[1]} colonnes")
    xls.close()
    
    # CEMAC
    cemac_path = OUT_DIR / "ICAE_CEMAC_OUT.xlsx"
    xls = pd.ExcelFile(cemac_path)
    print(f"  CEMAC ({cemac_path.name}):")
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        print(f"    {sheet}: {df.shape[0]} lignes x {df.shape[1]} colonnes")
    xls.close()
    
    print("  [OK] Tous les fichiers Excel sont lisibles")
except Exception as e:
    errors.append(f"Validation Excel: {e}")
    print(f"  [ERREUR] {e}")
    import traceback; traceback.print_exc()


# ====================================================================
# BILAN
# ====================================================================
print("\n" + "=" * 60)
if errors:
    print(f"BILAN : {len(errors)} ERREUR(S)")
    for e in errors:
        print(f"  - {e}")
else:
    print("BILAN : TOUS LES TESTS D'EXPORT REUSSIS !")
    print(f"  Fichiers generes dans : {OUT_DIR}")
    for f in sorted(OUT_DIR.iterdir()):
        print(f"    {f.name} ({f.stat().st_size:,} bytes)")
print("=" * 60)
