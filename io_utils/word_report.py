"""Génération de rapports Word (.docx) — Note ICAE et Note Nowcast."""
import io
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def _add_heading_beac(doc, text, level=1):
    """Ajoute un titre formaté BEAC."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"


def _add_body_text(doc, text):
    """Ajoute un paragraphe de texte corps."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _add_figure_caption(doc, caption):
    """Ajoute une légende de figure."""
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_image(doc, image_bytes, width=Cm(15)):
    """Ajoute une image depuis des bytes."""
    buf = io.BytesIO(image_bytes)
    doc.add_picture(buf, width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc, df: pd.DataFrame):
    """Ajoute un DataFrame comme table Word."""
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

    # Données
    for i, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            if isinstance(val, float):
                cell.text = f"{val:.2f}" if not np.isnan(val) else ""
            else:
                cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"

    return table


# ────────────────────────────────────────────────────────────────────────────
# Note ICAE CEMAC (format court, calqué sur le modèle existant)
# ────────────────────────────────────────────────────────────────────────────
def generate_note_icae(
    country_name: str,
    trimestre: str,
    annee: int,
    icae_data: dict,
    chart_evolution_bytes: bytes = None,
    chart_perspectives_bytes: bytes = None,
    logo_path: str = None,
    user_texts: dict = None,
) -> bytes:
    """
    Génère la Note ICAE (CEMAC ou pays).

    Parameters
    ----------
    user_texts : dict avec clés optionnelles :
        'accroche_evolution', 'paragraphs_evolution' (list[str]),
        'accroche_perspectives', 'paragraphs_perspectives' (list[str])
    """
    doc = Document()

    # Style par défaut
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    if user_texts is None:
        user_texts = {}

    # ── Logo ──────────────────────────────────────────────────────────────
    if logo_path and Path(logo_path).exists():
        doc.add_picture(str(logo_path), width=Cm(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Titre ─────────────────────────────────────────────────────────────
    title = (
        f"EVOLUTION DE L'INDICE COMPOSITE DES ACTIVITES ECONOMIQUES (ICAE) "
        f"DE {country_name.upper()} AU {trimestre} {annee} "
        f"ET PERSPECTIVES A COURT TERME"
    )
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # espace

    # ── Section 1 : Evolution récente ─────────────────────────────────────
    _add_heading_beac(doc, "Evolution récente")

    # Accroche
    accroche_evo = user_texts.get(
        "accroche_evolution",
        f"[Renseignez l'accroche : Le rythme de progression des activités "
        f"économiques a ... au {trimestre} {annee}.]"
    )
    p = doc.add_paragraph()
    run = p.add_run(accroche_evo)
    run.bold = True
    run.font.size = Pt(14)

    # Paragraphes d'analyse
    paras_evo = user_texts.get("paragraphs_evolution", [
        "[Paragraphe 1 : Dynamique globale du secteur productif. "
        "Utilisez les données et graphiques ci-contre pour rédiger.]",
        "[Paragraphe 2 : Activités extractives — pétrole, mines, etc.]",
        "[Paragraphe 3 : Industrie manufacturière, BTP, sylviculture.]",
        "[Paragraphe 4 : Transports, services et autres secteurs.]",
    ])
    for para in paras_evo:
        _add_body_text(doc, para)

    # Figure 1
    _add_figure_caption(
        doc,
        f"Figure 1 : Evolution de l'ICAE {country_name} en glissement annuel"
    )
    if chart_evolution_bytes:
        _add_image(doc, chart_evolution_bytes)
    else:
        _add_body_text(doc, "[Graphique non disponible — générez-le dans le module ICAE]")

    doc.add_paragraph()  # espace

    # ── Section 2 : Perspectives ──────────────────────────────────────────
    _add_heading_beac(doc, "Perspectives d'évolution de l'ICAE à court terme")

    accroche_persp = user_texts.get(
        "accroche_perspectives",
        f"[Renseignez l'accroche : ...le rythme de progression des activités "
        f"économiques devrait... au trimestre suivant.]"
    )
    p = doc.add_paragraph()
    run = p.add_run(accroche_persp)
    run.bold = True
    run.font.size = Pt(14)

    paras_persp = user_texts.get("paragraphs_perspectives", [
        "[Paragraphe 1 : Projections sectorielles.]",
        "[Paragraphe 2 : Facteurs de risque et incertitudes.]",
        "[Paragraphe 3 : Synthèse — chiffres clés de la projection.]",
    ])
    for para in paras_persp:
        _add_body_text(doc, para)

    # Figure 2
    _add_figure_caption(
        doc,
        f"Figure 2 : Evolution de l'ICAE {country_name} en glissement annuel "
        f"(avec projection)"
    )
    if chart_perspectives_bytes:
        _add_image(doc, chart_perspectives_bytes)
    else:
        _add_body_text(doc, "[Graphique non disponible — générez-le dans le module Prévisions]")

    # Sauvegarder en bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ────────────────────────────────────────────────────────────────────────────
# Note Nowcast (format technique)
# ────────────────────────────────────────────────────────────────────────────
import random as _random

# ── Banque de formulations pour la note Nowcast ───────────────────────────
_NW_INTRO = [
    ("La présente note expose les résultats des exercices de nowcasting du PIB "
     "réalisés dans le cadre du suivi conjoncturel infra-annuel. "
     "Le nowcasting consiste à estimer en temps réel le niveau de l'activité "
     "économique avant la publication officielle des comptes nationaux, en "
     "exploitant les indicateurs haute fréquence disponibles avec une moindre "
     "latence de publication."),
    ("Cette note technique présente l'estimation avancée (nowcast) du Produit "
     "Intérieur Brut (PIB) réalisée à partir d'indicateurs composites haute "
     "fréquence. La démarche vise à combler le délai habituel entre la fin "
     "d'un trimestre et la publication des comptes nationaux, afin de disposer "
     "d'une appréciation conjoncturelle rapide."),
    ("Les travaux de nowcasting présentés ci-après s'inscrivent dans la "
     "démarche de surveillance rapprochée du cycle économique. Ils permettent "
     "d'estimer l'évolution du PIB trimestriel à partir des indicateurs "
     "haute fréquence disponibles avant la clôture des comptes nationaux, "
     "réduisant ainsi significativement les délais d'information conjoncturelle."),
]

_NW_METH_INTRO = [
    ("Quatre modèles de nowcasting ont été mis en œuvre, couvrant différentes "
     "approches statistiques et économétriques afin de maximiser la robustesse "
     "des estimations."),
    ("L'exercice mobilise quatre approches complémentaires de nowcasting, "
     "dont les performances sont évaluées sur données rétrospectives. "
     "La confrontation des résultats renforce la fiabilité des estimations."),
    ("Les estimations reposent sur quatre méthodologies distinctes, "
     "sélectionnées pour leur complémentarité et leur robustesse "
     "vis-à-vis des spécificités structurelles des économies de la CEMAC."),
]

_NW_BRIDGE_DESC = (
    "**Bridge Equation** : Ce modèle établit une équation de pont entre les "
    "indicateurs haute fréquence agrégés en données trimestrielles et le PIB. "
    "Il s'appuie sur la première composante principale des indicateurs (PC1) "
    "ainsi que sur les deux indicateurs les mieux corrélés avec le PIB, "
    "introduits avec un retard d'un trimestre. L'estimation est réalisée par "
    "les moindres carrés ordinaires (MCO)."
)
_NW_UMIDAS_DESC = (
    "**U-MIDAS (Unrestricted Mixed-Data Sampling)** : Le modèle U-MIDAS "
    "mobilise directement le meilleur indicateur haute fréquence et ses deux "
    "retards comme régresseurs du PIB trimestriel, sans imposer de structure "
    "paramétrique de pondération sur les retards. Cette approche est "
    "particulièrement adaptée lorsqu'un seul indicateur domine dans la "
    "prédiction du PIB."
)
_NW_PC_DESC = (
    "**Régression sur composantes principales (PC)** : Cette méthode réduit "
    "la dimensionnalité de l'espace des indicateurs par Analyse en Composantes "
    "Principales (ACP), puis régresse le PIB sur les k premières composantes "
    "principales. Elle s'avère robuste en présence de multicolinéarité entre "
    "les indicateurs HF et permet d'exploiter l'information contenue dans "
    "l'ensemble du jeu de données."
)
_NW_DFM_DESC = (
    "**Dynamic Factor Model — DFM-lite** : Ce modèle à facteurs dynamiques "
    "extrait un facteur commun latent (PC1) à partir des indicateurs HF "
    "et modélise le PIB en fonction de ce facteur et de son retard. "
    "Il capture les co-mouvements cycliques entre l'ensemble des indicateurs "
    "et le PIB, tout en maintenant un nombre réduit de paramètres à estimer."
)

_NW_BEST_MODEL = [
    "Parmi les {n} modèles évalués, le modèle **{m}** s'est révélé le plus "
    "performant, avec un RMSE hors-échantillon de {rmse:.2f}.",
    "L'analyse comparative fait ressortir la supériorité du modèle **{m}** "
    "sur les données hors-échantillon (RMSE = {rmse:.2f}).",
    "Au terme de l'évaluation rétrospective, le modèle **{m}** est retenu "
    "comme référence principale, affichant le plus faible RMSE out-of-sample "
    "({rmse:.2f}).",
    "C'est le modèle **{m}** qui présente la meilleure capacité prédictive "
    "sur la période de validation hors-échantillon (RMSE = {rmse:.2f}).",
]

_NW_GROWTH_PHRASES = {
    "positive": [
        "Le nowcast {model} signale une **progression de l'activité économique** "
        "au {period}, estimée à {ga:+.1f}% en glissement annuel.",
        "Selon le modèle {model}, le PIB affiche une **croissance de {ga:.1f}%** "
        "en glissement annuel au {period}.",
        "Le nowcast indique une **expansion économique** au {period}, "
        "avec un glissement annuel estimé à +{ga:.1f}% par le modèle {model}.",
    ],
    "negative": [
        "Le nowcast {model} signale une **contraction de l'activité économique** "
        "au {period}, estimée à {ga:.1f}% en glissement annuel.",
        "Selon le modèle {model}, le PIB affiche un **recul de {ga:.1f}%** "
        "en glissement annuel au {period}.",
        "Le nowcast indique une **compression de l'activité** au {period}, "
        "avec un glissement annuel estimé à {ga:.1f}% par le modèle {model}.",
    ],
    "flat": [
        "Le nowcast {model} indique une **activité quasi-stable** au {period}, "
        "avec un glissement annuel proche de zéro ({ga:+.1f}%).",
        "Selon le modèle {model}, la croissance du PIB reste **modérée** "
        "au {period} ({ga:+.1f}% en glissement annuel).",
    ],
}

_NW_TREND_ACCEL = [
    "Cette dynamique s'inscrit dans un mouvement d'**accélération** par rapport "
    "au trimestre précédent, suggérant un raffermissement du cycle économique.",
    "La progression s'accélère par rapport au trimestre précédent, "
    "signe d'une **reprise progressive du dynamisme économique**.",
    "L'évolution est favorable et traduit une **montée en puissance** "
    "de l'activité par rapport aux trimestres antérieurs.",
]
_NW_TREND_SLOW = [
    "Néanmoins, cette progression marque un **ralentissement** par rapport "
    "au trimestre précédent, invitant à une vigilance maintenue.",
    "Le rythme de croissance s'infléchit par rapport au trimestre précédent, "
    "signalant un certain **essoufflement de la dynamique**.",
    "La croissance demeure positive mais **décélère**, ce qui appelle à "
    "un suivi attentif des indicateurs haute fréquence dans les prochains mois.",
]
_NW_TREND_STABLE = [
    "Cette performance s'inscrit dans la continuité des trimestres précédents, "
    "traduisant une **relative stabilité du rythme de croissance**.",
    "Le nowcast confirme la **persistance du régime conjoncturel** actuel, "
    "sans rupture notable par rapport aux dernières estimations.",
]

_NW_CONCLUSION_VARIANTS = [
    ("En synthèse, les résultats des modèles de nowcasting convergent vers une "
     "appréciation cohérente de la conjoncture, le modèle {best_model} ayant "
     "démontré les meilleures performances prédictives. Ces estimations "
     "constituent un signal avancé utile pour l'orientation des analyses "
     "conjoncturelles, dans l'attente des comptes nationaux définitifs. "
     "Il convient toutefois de les interpréter avec précaution, la qualité "
     "des nowcasts demeurant conditionnée par la disponibilité et la "
     "représentativité des indicateurs haute fréquence mobilisés."),
    ("Les exercices de nowcasting présentés dans cette note offrent une "
     "estimation rapide de l'évolution du PIB, comblant en partie le "
     "délai d'information conjoncturelle. Les résultats obtenus, bien que "
     "fondés sur des modèles éprouvés, restent tributaires des révisions "
     "ultérieures des indicateurs haute fréquence et des comptes nationaux. "
     "Le modèle {best_model} est recommandé à titre principal pour les "
     "prochains exercices, sous réserve de réévaluation périodique "
     "des performances hors-échantillon."),
    ("L'analyse comparative des performances démontre l'intérêt d'une "
     "approche multi-modèles pour le nowcasting du PIB. Le modèle {best_model}, "
     "sélectionné sur la base du RMSE hors-échantillon, servira de référence "
     "principale pour les estimations à venir. Les travaux seront poursuivis "
     "par l'actualisation régulière des indicateurs HF et la révision "
     "périodique du portefeuille de modèles, en particulier lors des "
     "changements structurels majeurs dans les économies concernées."),
]


def _nw_ga_period_label(idx) -> str:
    """Retourne un label lisible pour le dernier trimestre du forecast."""
    try:
        if hasattr(idx, 'year') and hasattr(idx, 'quarter'):
            mois = {1: "T1", 2: "T2", 3: "T3", 4: "T4"}
            return f"{mois[idx.quarter]} {idx.year}"
        return str(idx)
    except Exception:
        return str(idx)


def _nw_compute_ga(fc_series: "pd.Series") -> tuple:
    """Retourne (dernier trimestre, GA yoy) depuis une série forecast."""
    try:
        s = fc_series.dropna()
        if len(s) < 5:
            return None, None
        last_val = float(s.iloc[-1])
        year_ago = float(s.iloc[-5]) if len(s) >= 5 else None
        if year_ago is not None and year_ago != 0:
            ga = (last_val / year_ago - 1) * 100
        else:
            ga = None
        last_idx = s.index[-1]
        return _nw_ga_period_label(last_idx), ga
    except Exception:
        return None, None


def generate_note_nowcast(
    results_by_country: dict,
    chart_bytes: dict = None,
    logo_path: str = None,
    params: dict = None,
) -> bytes:
    """Génère la Note Nowcast technique avec template conditionnel enrichi.

    Parameters
    ----------
    results_by_country : dict
        {code: {metrics_df, pib_q, forecasts, best_model, country_name}}
    params : dict optionnel
        {periode, models_used, hf_vars}
    """
    if chart_bytes is None:
        chart_bytes = {}
    if params is None:
        params = {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Logo ──────────────────────────────────────────────────────────────
    if logo_path and Path(logo_path).exists():
        doc.add_picture(str(logo_path), width=Cm(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Titre ─────────────────────────────────────────────────────────────
    periode_label = params.get("periode", pd.Timestamp.now().strftime("%B %Y"))
    p = doc.add_paragraph()
    run = p.add_run(
        f"NOTE TECHNIQUE SUR L'ESTIMATION EN TEMPS RÉEL (NOWCAST) "
        f"DU PRODUIT INTÉRIEUR BRUT — {periode_label.upper()}"
    )
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 1. Introduction ───────────────────────────────────────────────────
    _add_heading_beac(doc, "I. Introduction et objectifs")
    _add_body_text(doc, _random.choice(_NW_INTRO))
    _add_body_text(doc,
        "Les indicateurs haute fréquence mobilisés dans cet exercice sont "
        "agrégés en données trimestrielles selon leur nature économique : "
        "somme pour les flux (production, exportations), moyenne pour les "
        "taux et indices (IPC, taux d'intérêt), et valeur de fin de période "
        "pour les stocks (masse monétaire, crédit à l'économie). "
        "Le PIB annuel, lorsqu'il n'est pas disponible en fréquence "
        "trimestrielle, est désagrégé par la méthode de Chow-Lin ou "
        "Denton-Cholette selon la disponibilité d'un indicateur de référence."
    )

    # ── 2. Indicateurs utilisés ───────────────────────────────────────────
    hf_vars = params.get("hf_vars", [])
    _add_heading_beac(doc, "II. Indicateurs haute fréquence utilisés")
    if hf_vars:
        _add_body_text(doc,
            f"L'estimation s'appuie sur {len(hf_vars)} indicateur(s) haute "
            "fréquence sélectionné(s) selon les critères de disponibilité, "
            "de pertinence économique et de corrélation avec le PIB :"
        )
        for v in hf_vars:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(v)).font.size = Pt(11)
    else:
        _add_body_text(doc,
            "Les indicateurs haute fréquence ont été sélectionnés depuis la "
            "feuille de codification du classeur ICAE (variables avec "
            "PRIOR > 0 et statut Actif). La liste complète est disponible "
            "dans le fichier Excel exporté depuis le Module 3."
        )

    # ── 3. Méthodes ───────────────────────────────────────────────────────
    _add_heading_beac(doc, "III. Méthodes d'estimation")
    _add_body_text(doc, _random.choice(_NW_METH_INTRO))
    doc.add_paragraph()

    models_used = params.get("models_used", ["Bridge", "U-MIDAS", "PC", "DFM"])
    method_map = {
        "Bridge": _NW_BRIDGE_DESC,
        "U-MIDAS": _NW_UMIDAS_DESC,
        "PC": _NW_PC_DESC,
        "DFM": _NW_DFM_DESC,
    }
    for model_key in ["Bridge", "U-MIDAS", "PC", "DFM"]:
        if model_key in models_used or not models_used:
            desc = method_map.get(model_key, f"[{model_key} — description non disponible]")
            p = doc.add_paragraph()
            # Titre de méthode en gras (extrait jusqu'au premier ':')
            parts = desc.split("**", 2)
            if len(parts) >= 3:
                run_b = p.add_run(parts[1])
                run_b.bold = True
                run_b.font.size = Pt(12)
                run_b.font.name = "Times New Roman"
                run_r = p.add_run(parts[2].lstrip(":").lstrip(" ") if "**" in parts[2]
                                  else parts[2])
                run_r.font.size = Pt(12)
                run_r.font.name = "Times New Roman"
            else:
                p.add_run(desc).font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── 4. Résultats par pays ─────────────────────────────────────────────
    _add_heading_beac(doc, "IV. Résultats par pays")

    best_models_global = []
    for c_idx, (code, c_results) in enumerate(results_by_country.items(), start=1):
        country_name = c_results.get("country_name", code)
        best_model = c_results.get("best_model")
        pib_q = c_results.get("pib_q")
        forecasts = c_results.get("forecasts", {})
        metrics_df = c_results.get("metrics_df", pd.DataFrame())

        # Sous-titre pays
        p = doc.add_paragraph()
        run = p.add_run(f"{c_idx}. {country_name}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"
        doc.add_paragraph()

        # 4.a Tableau de performance
        if not metrics_df.empty:
            _add_body_text(doc,
                "Le tableau ci-dessous présente les métriques de performance "
                "des modèles évalués, calculées sur les données in-sample "
                "(ajustement global) et out-of-sample (validation rétrospective) :"
            )
            _add_table(doc, metrics_df)
            doc.add_paragraph()

        # 4.b Sélection du meilleur modèle
        if best_model and not metrics_df.empty:
            _model_col = next(
                (c for c in metrics_df.columns if "mod" in str(c).lower()), None
            )
            best_row = (metrics_df[metrics_df[_model_col] == best_model]
                        if _model_col else pd.DataFrame())
            rmse_out = (best_row["RMSE (out)"].values[0]
                        if len(best_row) > 0 and "RMSE (out)" in best_row.columns
                        else np.nan)
            best_models_global.append(best_model)
            phrase = _random.choice(_NW_BEST_MODEL).format(
                n=len(metrics_df),
                m=best_model,
                rmse=rmse_out if not np.isnan(rmse_out) else 0,
            )
            _add_body_text(doc, phrase)

        # 4.c Analyse du nowcast (GA et scénario)
        best_fc = forecasts.get(best_model) if best_model else None
        if best_fc is not None and len(best_fc.dropna()) > 0:
            period_label, ga = _nw_compute_ga(best_fc)
            if period_label and ga is not None:
                if ga > 0.5:
                    key = "positive"
                elif ga < -0.5:
                    key = "negative"
                else:
                    key = "flat"
                ga_phrase = _random.choice(_NW_GROWTH_PHRASES[key]).format(
                    model=best_model, period=period_label, ga=ga
                )
                _add_body_text(doc, ga_phrase)

                # Tendance (comparer au trimestre précédent)
                s = best_fc.dropna()
                if len(s) >= 6:
                    ga_prev = None
                    try:
                        ga_prev = (float(s.iloc[-2]) / float(s.iloc[-6]) - 1) * 100
                    except Exception:
                        pass
                    if ga_prev is not None:
                        if ga > ga_prev + 0.3 and ga > 0:
                            _add_body_text(doc, _random.choice(_NW_TREND_ACCEL))
                        elif ga < ga_prev - 0.3:
                            _add_body_text(doc, _random.choice(_NW_TREND_SLOW))
                        else:
                            _add_body_text(doc, _random.choice(_NW_TREND_STABLE))

            # Niveaux observés vs nowcastés
            if pib_q is not None and len(pib_q.dropna()) > 0:
                last_obs = float(pib_q.dropna().iloc[-1])
                last_obs_period = _nw_ga_period_label(pib_q.dropna().index[-1])
                last_nowcast = float(best_fc.dropna().iloc[-1])
                last_nowcast_period = _nw_ga_period_label(best_fc.dropna().index[-1])
                _add_body_text(doc,
                    f"À titre indicatif, le PIB observé s'établissait à "
                    f"{last_obs:,.1f} au {last_obs_period}. "
                    f"Le nowcast {best_model} pour le {last_nowcast_period} "
                    f"est estimé à {last_nowcast:,.1f}."
                )

        # 4.d Graphique
        if code in chart_bytes:
            _add_figure_caption(
                doc,
                f"Figure {c_idx} : PIB observé et Nowcasts — {country_name}"
            )
            _add_image(doc, chart_bytes[code])
        else:
            _add_body_text(doc,
                "[Graphique non disponible — lancez le Nowcast dans le Module 3 "
                "et générez le rapport depuis ce module.]"
            )

        doc.add_paragraph()  # espace entre pays

    # ── 5. Synthèse et conclusion ─────────────────────────────────────────
    _add_heading_beac(doc, "V. Synthèse et recommandations")
    from collections import Counter
    if best_models_global:
        dominant = Counter(best_models_global).most_common(1)[0][0]
    else:
        dominant = "Bridge"
    _add_body_text(doc, _random.choice(_NW_CONCLUSION_VARIANTS).format(
        best_model=dominant))
    _add_body_text(doc,
        "Il est recommandé de renouveler cet exercice à chaque actualisation "
        "trimestrielle des indicateurs haute fréquence, et de procéder à une "
        "révision complète du portefeuille d'indicateurs tous les deux à trois "
        "ans, en tenant compte des évolutions structurelles des économies "
        "et des nouvelles sources de données disponibles."
    )

    # ── Pied de note ──────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        f"Note générée le {pd.Timestamp.now().strftime('%d %B %Y')} "
        f"par l'application ICAE & PIB Nowcast — BEAC/DCCE."
    )
    run.font.size = Pt(10)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
